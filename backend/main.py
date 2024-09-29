
# # # import hnswlib
# # # import numpy as np
# # # from langchain.embeddings import OpenAIEmbeddings
# # # from langchain.schema import Document
# # # from dotenv import load_dotenv
# # # import os

# # # # Load environment variables from the .env file
# # # base_dir = os.path.dirname(os.path.abspath(__file__))  # Get the directory of the current script
# # # env_path = os.path.join(base_dir, '../config/.env')    # Construct the full path to the .env file
# # # load_dotenv(dotenv_path=env_path)                      # Load the .env file

# # # # Get the OpenAI API key from environment variables
# # # openai_api_key = os.getenv("OPENAI_API_KEY")
# # # if not openai_api_key:
# # #     raise ValueError("API Key not found in environment variables. Please check your .env file.")

# # # # Initialize OpenAI embeddings with the API key
# # # embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)

# # # # Example: Add some sample documents (In a real scenario, load a larger dataset)
# # # documents = [
# # #     Document(page_content="HackGT is an annual hackathon at Georgia Tech."),
# # #     Document(page_content="LangChain is a framework for building applications with language models."),
# # #     Document(page_content="Georgia Tech is a public research university in Atlanta, Georgia."),
# # #     # Add more documents here for larger datasets
# # # ]

# # # # Extract document texts
# # # texts = [doc.page_content for doc in documents]

# # # # Generate embeddings for the documents
# # # doc_embeddings = embeddings.embed_documents(texts)
# # # doc_embeddings = np.array(doc_embeddings)

# # # # Number of dimensions for each vector
# # # embedding_dim = doc_embeddings.shape[1]

# # # # Initialize the HNSW index
# # # hnsw_index = hnswlib.Index(space='cosine', dim=embedding_dim)  # 'cosine' or 'l2' distance

# # # # Initialize the index with the expected number of elements
# # # num_elements = len(doc_embeddings)
# # # hnsw_index.init_index(max_elements=num_elements, ef_construction=100, M=16)

# # # # Add documents to the HNSW index
# # # hnsw_index.add_items(doc_embeddings, list(range(num_elements)))

# # # # Set parameters for search speed vs accuracy
# # # hnsw_index.set_ef(10)  # Higher value means better accuracy, but slower search

# # # print("Documents added to HNSW index successfully.")

# # # # Define a minimal test query function
# # # def test_similarity_search(query, top_k=2):
# # #     # Get query embedding
# # #     query_embedding = embeddings.embed_documents([query])[0]
    
# # #     # Perform a query search in HNSWlib
# # #     labels, distances = hnsw_index.knn_query(np.array([query_embedding]), k=top_k)
    
# # #     # Print the closest documents
# # #     print(f"Query: {query}")
# # #     print("Closest Documents:")
# # #     for label, distance in zip(labels[0], distances[0]):
# # #         print(f"- {documents[label].page_content} (distance: {distance})")

# # # # Test the similarity search
# # # if __name__ == '__main__':
# # #     test_similarity_search("What is HackGT?")
# # #     test_similarity_search("What is LangChain?")
# # import os
# # import hnswlib
# # import numpy as np
# # import pandas as pd
# # from langchain.embeddings import OpenAIEmbeddings
# # from langchain.schema import Document
# # from dotenv import load_dotenv
# # from docx import Document as DocxDocument
# # from PyPDF2 import PdfReader
# # import docx2txt

# # # Load environment variables from the .env file
# # base_dir = os.path.dirname(os.path.abspath(__file__))  # Get the directory of the current script
# # env_path = os.path.join(base_dir, '../config/.env')    # Construct the full path to the .env file
# # load_dotenv(dotenv_path=env_path)                      # Load the .env file

# # # Get the OpenAI API key from environment variables
# # openai_api_key = os.getenv("OPENAI_API_KEY")
# # if not openai_api_key:
# #     raise ValueError("API Key not found in environment variables. Please check your .env file.")

# # # Initialize OpenAI embeddings with the API key
# # embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)

# # # Initialize HNSWlib index with higher values for M and ef_construction
# # embedding_dim = 1536  # Number of dimensions for OpenAI embeddings
# # num_elements = 1000   # Placeholder value, update with actual number of documents
# # M = 96                # Increased M value for better connectivity
# # ef_construction = 400  # Increased ef_construction value for better graph quality
# # hnsw_index = hnswlib.Index(space='cosine', dim=embedding_dim)
# # hnsw_index.init_index(max_elements=num_elements, ef_construction=ef_construction, M=M)
# # print("HNSW index initialized.")

# # # List to store document objects
# # documents = []

# # # Function to extract text from .txt files
# # def extract_text_from_txt(file_path):
# #     try:
# #         with open(file_path, 'r', encoding='utf-8') as file:
# #             return file.read()
# #     except Exception as e:
# #         print(f"Error extracting {file_path} as .txt: {e}")
# #         return ""

# # # Function to extract text from .doc files using docx2txt
# # def extract_text_from_doc(file_path):
# #     try:
# #         text = docx2txt.process(file_path)
# #     except Exception as e:
# #         print(f"Error extracting {file_path} with docx2txt: {e}")
# #         text = ""
# #     return text

# # # Function to extract text from .docx files
# # def extract_text_from_docx(file_path):
# #     try:
# #         doc = DocxDocument(file_path)
# #         return "\n".join([para.text for para in doc.paragraphs])
# #     except Exception as e:
# #         print(f"Error extracting {file_path} with python-docx: {e}")
# #         return ""

# # # Function to extract text from .pdf files
# # def extract_text_from_pdf(file_path):
# #     try:
# #         text = ""
# #         with open(file_path, 'rb') as file:
# #             reader = PdfReader(file)
# #             for page in reader.pages:
# #                 text += page.extract_text() or ""
# #     except Exception as e:
# #         print(f"Error extracting {file_path} with PyPDF2: {e}")
# #         text = ""
# #     return text

# # # Function to extract text from .xlsx files using pandas
# # def extract_text_from_xlsx(file_path):
# #     try:
# #         df = pd.read_excel(file_path, dtype=str)  # Read entire sheet as text
# #         return "\n".join(df.fillna("").apply(lambda x: " ".join(x), axis=1))
# #     except Exception as e:
# #         print(f"Error extracting {file_path} with pandas: {e}")
# #         return ""

# # # Function to extract text from .csv files using pandas
# # def extract_text_from_csv(file_path):
# #     try:
# #         df = pd.read_csv(file_path, dtype=str)  # Read entire CSV as text
# #         return "\n".join(df.fillna("").apply(lambda x: " ".join(x), axis=1))
# #     except Exception as e:
# #         print(f"Error extracting {file_path} with pandas: {e}")
# #         return ""

# # # Define a function to load documents from a folder and extract text
# # def load_documents_from_folder(folder_path):
# #     for filename in os.listdir(folder_path):
# #         file_path = os.path.join(folder_path, filename)
# #         if filename.endswith(".txt"):
# #             content = extract_text_from_txt(file_path)
# #         elif filename.endswith(".doc"):
# #             content = extract_text_from_doc(file_path)
# #         elif filename.endswith(".docx"):
# #             content = extract_text_from_docx(file_path)
# #         elif filename.endswith(".pdf"):
# #             content = extract_text_from_pdf(file_path)
# #         elif filename.endswith(".xlsx"):
# #             content = extract_text_from_xlsx(file_path)
# #         elif filename.endswith(".csv"):
# #             content = extract_text_from_csv(file_path)
# #         else:
# #             print(f"Unsupported file type: {filename}")
# #             continue
        
# #         if content.strip():
# #             documents.append(Document(page_content=content))
# #             print(f"Loaded: {filename}")

# # # Path to the folder containing various files
# # folder_path = os.path.join(base_dir, '../data')  # Update this path as needed

# # # Load documents from the folder
# # load_documents_from_folder(folder_path)

# # # Embed the loaded documents and add them to the index
# # if documents:
# #     print(f"Loaded {len(documents)} documents.")
# #     texts = [doc.page_content for doc in documents]
# #     doc_embeddings = embeddings.embed_documents(texts)
# #     doc_embeddings = np.array(doc_embeddings)
    
# #     for i, embedding in enumerate(doc_embeddings):
# #         hnsw_index.add_items([embedding], [i])
    
# #     hnsw_index.set_ef(100)  # Set a larger ef for querying
# #     print("Documents added to HNSW index successfully.")
# # else:
# #     print("No documents found in the specified folder.")

# # # Function to search the index with a query
# # def search_index(query, top_k=3):  # Reduce top_k temporarily to avoid large queries
# #     query_embedding = embeddings.embed_documents([query])[0]
# #     hnsw_index.set_ef(max(300, top_k * 100))  # Set ef for querying, ef should be much larger
# #     labels, distances = hnsw_index.knn_query(np.array([query_embedding]), k=top_k)
# #     results = [{"text": documents[label].page_content, "distance": distance}
# #                for label, distance in zip(labels[0], distances[0])]
# #     return results

# # # Test the similarity search with some queries
# # if __name__ == '__main__':
# #     query1 = "What is HackGT?"
# #     query2 = "What is LangChain?"
# #     query3 = "Where is Georgia Tech located?"
# #     query4 = "What is FAISS used for?"

# #     response1 = search_index(query1)
# #     response2 = search_index(query2)
# #     response3 = search_index(query3)
# #     response4 = search_index(query4)

# #     print("Response to Query 1:", response1)
# #     print("Response to Query 2:", response2)
# #     print("Response to Query 3:", response3)
# #     print("Response to Query 4:", response4)


# import os
# import hnswlib
# import numpy as np
# import pandas as pd
# from langchain.embeddings import OpenAIEmbeddings
# from langchain.schema import Document
# from dotenv import load_dotenv
# from docx import Document as DocxDocument
# from PyPDF2 import PdfReader
# import docx2txt

# # Load environment variables from the .env file
# base_dir = os.path.dirname(os.path.abspath(__file__))  # Get the directory of the current script
# env_path = os.path.join(base_dir, '../config/.env')    # Construct the full path to the .env file
# load_dotenv(dotenv_path=env_path)                      # Load the .env file

# # Get the OpenAI API key from environment variables
# openai_api_key = os.getenv("OPENAI_API_KEY")
# if not openai_api_key:
#     raise ValueError("API Key not found in environment variables. Please check your .env file.")

# # Initialize OpenAI embeddings with the API key
# embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)

# # Initialize HNSWlib index with higher values for M and ef_construction
# embedding_dim = 1536  # Number of dimensions for OpenAI embeddings
# num_elements = 1000   # Placeholder value, update with actual number of documents
# M = 96                # Increased M value for better connectivity
# ef_construction = 400  # Increased ef_construction value for better graph quality
# hnsw_index = hnswlib.Index(space='cosine', dim=embedding_dim)
# hnsw_index.init_index(max_elements=num_elements, ef_construction=ef_construction, M=M)
# print("HNSW index initialized.")

# # List to store document objects
# documents = []

# # Function to extract text from .txt files
# def extract_text_from_txt(file_path):
#     try:
#         with open(file_path, 'r', encoding='utf-8') as file:
#             return file.read()
#     except Exception as e:
#         print(f"Error extracting {file_path} as .txt: {e}")
#         return ""

# # Function to extract text from .doc files using docx2txt
# def extract_text_from_doc(file_path):
#     try:
#         text = docx2txt.process(file_path)
#     except Exception as e:
#         print(f"Error extracting {file_path} with docx2txt: {e}")
#         text = ""
#     return text

# # Function to extract text from .docx files
# def extract_text_from_docx(file_path):
#     try:
#         doc = DocxDocument(file_path)
#         return "\n".join([para.text for para in doc.paragraphs])
#     except Exception as e:
#         print(f"Error extracting {file_path} with python-docx: {e}")
#         return ""

# # Function to extract text from .pdf files
# def extract_text_from_pdf(file_path):
#     try:
#         text = ""
#         with open(file_path, 'rb') as file:
#             reader = PdfReader(file)
#             for page in reader.pages:
#                 text += page.extract_text() or ""
#     except Exception as e:
#         print(f"Error extracting {file_path} with PyPDF2: {e}")
#         text = ""
#     return text

# # Function to extract text from .xlsx files using pandas
# def extract_text_from_xlsx(file_path):
#     try:
#         df = pd.read_excel(file_path, dtype=str)  # Read entire sheet as text
#         return "\n".join(df.fillna("").apply(lambda x: " ".join(x), axis=1))
#     except Exception as e:
#         print(f"Error extracting {file_path} with pandas: {e}")
#         return ""

# # Function to extract text from .csv files using pandas
# def extract_text_from_csv(file_path):
#     try:
#         df = pd.read_csv(file_path, dtype=str)  # Read entire CSV as text
#         return "\n".join(df.fillna("").apply(lambda x: " ".join(x), axis=1))
#     except Exception as e:
#         print(f"Error extracting {file_path} with pandas: {e}")
#         return ""

# # Define a function to load documents from a folder and extract text
# def load_documents_from_folder(folder_path):
#     for filename in os.listdir(folder_path):
#         file_path = os.path.join(folder_path, filename)
#         if filename.endswith(".txt"):
#             content = extract_text_from_txt(file_path)
#         elif filename.endswith(".doc"):
#             content = extract_text_from_doc(file_path)
#         elif filename.endswith(".docx"):
#             content = extract_text_from_docx(file_path)
#         elif filename.endswith(".pdf"):
#             content = extract_text_from_pdf(file_path)
#         elif filename.endswith(".xlsx"):
#             content = extract_text_from_xlsx(file_path)
#         elif filename.endswith(".csv"):
#             content = extract_text_from_csv(file_path)
#         else:
#             print(f"Unsupported file type: {filename}")
#             continue
        
#         if content.strip():
#             documents.append(Document(page_content=content))
#             print(f"Loaded: {filename}")

# # Path to the folder containing various files
# folder_path = os.path.join(base_dir, '../data')  # Update this path as needed

# # Load documents from the folder
# load_documents_from_folder(folder_path)

# # Embed the loaded documents and add them to the index
# if documents:
#     print(f"Loaded {len(documents)} documents.")
#     texts = [doc.page_content for doc in documents]
#     doc_embeddings = embeddings.embed_documents(texts)
#     doc_embeddings = np.array(doc_embeddings)
    
#     for i, embedding in enumerate(doc_embeddings):
#         hnsw_index.add_items([embedding], [i])
    
#     hnsw_index.set_ef(300)  # Set a larger ef for querying
#     print("Documents added to HNSW index successfully.")
# else:
#     print("No documents found in the specified folder.")

# # Function to search the index with a query
# def search_index(query, top_k=5):  # Increase top_k and apply filter
#     query_embedding = embeddings.embed_documents([query])[0]
#     hnsw_index.set_ef(max(300, top_k * 2))  # Set ef for querying, ef should be much larger
#     labels, distances = hnsw_index.knn_query(np.array([query_embedding]), k=top_k)
#     results = [{"text": documents[label].page_content, "distance": distance}
#                for label, distance in zip(labels[0], distances[0])]
    
#     # Filter results with a distance threshold
#     threshold = 0.15  # Only include documents closer than this distance
#     filtered_results = [result for result in results if result["distance"] <= threshold]
#     return filtered_results if filtered_results else results  # Return filtered results or all if none match

# # Test the similarity search with some queries
# if __name__ == '__main__':
#     query1 = "What is HackGT?"
#     query2 = "What is LangChain?"
#     query3 = "Where is Georgia Tech located?"
#     query4 = "What is FAISS used for?"

#     response1 = search_index(query1)
#     response2 = search_index(query2)
#     response3 = search_index(query3)
#     response4 = search_index(query4)

#     print("Response to Query 1:", response1)
#     print("Response to Query 2:", response2)
#     print("Response to Query 3:", response3)
#     print("Response to Query 4:", response4)


import os
import hnswlib
import numpy as np
import pandas as pd
from langchain.embeddings import OpenAIEmbeddings
from langchain.schema import Document
from dotenv import load_dotenv
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import docx2txt

# Load environment variables from the .env file
base_dir = os.path.dirname(os.path.abspath(__file__))
env_path = os.path.join(base_dir, '../config/.env')
load_dotenv(dotenv_path=env_path)

# Get the OpenAI API key from environment variables
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    raise ValueError("API Key not found in environment variables. Please check your .env file.")

# Initialize OpenAI embeddings with the API key
embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)

# Initialize HNSWlib index
embedding_dim = 1536
num_elements = 1000  # Placeholder value, update with actual number of documents
M = 64
ef_construction = 400
hnsw_index = hnswlib.Index(space='cosine', dim=embedding_dim)
hnsw_index.init_index(max_elements=num_elements, ef_construction=ef_construction, M=M)
print("HNSW index initialized.")

# List to store document objects
documents = []

# Function to extract text from various file types
def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_text_from_doc(file_path):
    return docx2txt.process(file_path)

def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        reader = PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_xlsx(file_path):
    df = pd.read_excel(file_path, dtype=str)
    return "\n".join(df.fillna("").apply(lambda x: " ".join(x), axis=1))

def extract_text_from_csv(file_path):
    df = pd.read_csv(file_path, dtype=str)
    return "\n".join(df.fillna("").apply(lambda x: " ".join(x), axis=1))

# Load documents from a folder and extract text
def load_documents_from_folder(folder_path):
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(".txt"):
            content = extract_text_from_txt(file_path)
        elif filename.endswith(".doc"):
            content = extract_text_from_doc(file_path)
        elif filename.endswith(".docx"):
            content = extract_text_from_docx(file_path)
        elif filename.endswith(".pdf"):
            content = extract_text_from_pdf(file_path)
        elif filename.endswith(".xlsx"):
            content = extract_text_from_xlsx(file_path)
        elif filename.endswith(".csv"):
            content = extract_text_from_csv(file_path)
        else:
            print(f"Unsupported file type: {filename}")
            continue
        
        if content.strip():
            documents.append(Document(page_content=content))
            print(f"Loaded: {filename}")

# Load documents from the specified folder
folder_path = os.path.join(base_dir, '../data')
load_documents_from_folder(folder_path)

# Embed the loaded documents and add them to the index
if documents:
    print(f"Loaded {len(documents)} documents.")
    texts = [doc.page_content for doc in documents]
    doc_embeddings = embeddings.embed_documents(texts)
    doc_embeddings = np.array(doc_embeddings)
    
    for i, embedding in enumerate(doc_embeddings):
        hnsw_index.add_items([embedding], [i])
    
    hnsw_index.set_ef(100)
    print("Documents added to HNSW index successfully.")
else:
    print("No documents found in the specified folder.")

# Function to search the index with a query
def search_index(query, top_k=3):
    query_embedding = embeddings.embed_documents([query])[0]
    hnsw_index.set_ef(max(300, top_k * 100))  
    labels, distances = hnsw_index.knn_query(np.array([query_embedding]), k=top_k)
    
    results = [{"text": documents[label].page_content, "distance": distance}
               for label, distance in zip(labels[0], distances[0])]
    return results

# Define a function to enhance the assistant's response
def generate_response(query):
    results = search_index(query)
    
    # Construct a conversational response
    if results:
        response_text = f"I found some information for you:\n"
        for result in results:
            response_text += f"- {result['text']} (similarity: {1 - result['distance']:.2f})\n"
        return response_text
    else:
        return "I'm sorry, I couldn't find any relevant information."

# Test the search functionality
if __name__ == '__main__':
    query1 = "What is HackGT?"
    query2 = "What is LangChain?"
    query3 = "Where is Georgia Tech located?"
    query4 = "What is FAISS used for?"

    print("Response to Query 1:", generate_response(query1))
    print("Response to Query 2:", generate_response(query2))
    print("Response to Query 3:", generate_response(query3))
    print("Response to Query 4:", generate_response(query4))
