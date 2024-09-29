import os
import hnswlib
import numpy as np
import pandas as pd
import pickle
import time
from langchain.embeddings import OpenAIEmbeddings
from langchain.schema import Document
from dotenv import load_dotenv
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import docx2txt
from langchain.text_splitter import RecursiveCharacterTextSplitter
import openai
from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename

# ==========================
# Configuration and Setup
# ==========================

# Load environment variables from the .env file
base_dir = os.path.dirname(os.path.abspath(__file__))
env_path = os.path.join(base_dir, '../config/.env')
load_dotenv(dotenv_path=env_path)

# Get the OpenAI API key from environment variables
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    raise ValueError("API Key not found in environment variables. Please check your .env file.")

# Initialize OpenAI embeddings with the API key
embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)

# Initialize OpenAI for ChatCompletion
openai.api_key = openai_api_key

# Initialize Flask app
app = Flask(__name__, template_folder='../frontend/templates', static_folder='../frontend/static')

# Set up the data folder for uploads
UPLOAD_FOLDER = os.path.join(base_dir, '../data')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Allowable file extensions for upload
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx', 'xlsx', 'csv'}

# Ensure the upload folder exists
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# ==========================
# Helper Functions
# ==========================

def allowed_file(filename):
    """Check if the file has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ==========================
# Indexer Class
# ==========================

class Indexer:
    def __init__(self, embedding_dim=1536, num_elements=100000, M=64, ef_construction=400, ef=300):
        self.embedding_dim = embedding_dim
        self.num_elements = num_elements
        self.M = M
        self.ef_construction = ef_construction
        self.ef = ef
        self.hnsw_index = hnswlib.Index(space='cosine', dim=self.embedding_dim)
        self.hnsw_index.init_index(max_elements=self.num_elements, ef_construction=self.ef_construction, M=self.M)
        self.hnsw_index.set_ef(self.ef)
        self.documents = []
        print("HNSW index initialized.")

    def extract_text(self, file_path):
        """Extract text based on file extension."""
        if file_path.endswith(".txt"):
            return self.extract_text_from_txt(file_path)
        elif file_path.endswith(".doc"):
            return self.extract_text_from_doc(file_path)
        elif file_path.endswith(".docx"):
            return self.extract_text_from_docx(file_path)
        elif file_path.endswith(".pdf"):
            return self.extract_text_from_pdf(file_path)
        elif file_path.endswith(".xlsx"):
            return self.extract_text_from_xlsx(file_path)
        elif file_path.endswith(".csv"):
            return self.extract_text_from_csv(file_path)
        else:
            print(f"Unsupported file type: {file_path}")
            return ""

    def extract_text_from_txt(self, file_path):
        """Extract text from a .txt file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting {file_path} as .txt: {e}")
            return ""

    def extract_text_from_doc(self, file_path):
        """Extract text from a .doc file using docx2txt."""
        try:
            text = docx2txt.process(file_path)
            return text
        except Exception as e:
            print(f"Error extracting {file_path} with docx2txt: {e}")
            return ""

    def extract_text_from_docx(self, file_path):
        """Extract text from a .docx file."""
        try:
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"Error extracting {file_path} with python-docx: {e}")
            return ""

    def extract_text_from_pdf(self, file_path):
        """Extract text from a .pdf file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() or ""
            return text
        except Exception as e:
            print(f"Error extracting {file_path} with PyPDF2: {e}")
            return ""

    def extract_text_from_xlsx(self, file_path):
        """Extract text from a .xlsx file using pandas."""
        try:
            df = pd.read_excel(file_path, dtype=str)
            return "\n".join(df.fillna("").apply(lambda x: " ".join(x), axis=1))
        except Exception as e:
            print(f"Error extracting {file_path} with pandas: {e}")
            return ""

    def extract_text_from_csv(self, file_path):
        """Extract text from a .csv file using pandas."""
        try:
            df = pd.read_csv(file_path, dtype=str)
            return "\n".join(df.fillna("").apply(lambda x: " ".join(x), axis=1))
        except Exception as e:
            print(f"Error extracting {file_path} with pandas: {e}")
            return ""

    def load_documents_from_folder(self, folder_path):
        """Load and split documents from a specified folder."""
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            content = self.extract_text(file_path)
            if content.strip():
                # Split the content into chunks
                chunks = text_splitter.split_text(content)
                for chunk in chunks:
                    self.documents.append(Document(page_content=chunk))
                    print(f"Loaded chunk from: {filename}")

    def load_and_index_file(self, file_path):
        """Load and index a single file."""
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        content = self.extract_text(file_path)
        if content.strip():
            # Split the content into chunks
            chunks = text_splitter.split_text(content)
            new_docs = []
            for chunk in chunks:
                new_doc = Document(page_content=chunk)
                self.documents.append(new_doc)
                new_docs.append(chunk)
                print(f"Loaded chunk from: {os.path.basename(file_path)}")

            # Embed the new chunks
            texts = new_docs
            batch_size = 100  # Adjust as needed
            doc_embeddings = []
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i+batch_size]
                embeddings_batch = embeddings.embed_documents(batch_texts)
                doc_embeddings.extend(embeddings_batch)
                print(f"Embedded batch {i//batch_size +1}")

            doc_embeddings = np.array(doc_embeddings)

            # Add embeddings to HNSW index
            start_idx = len(self.documents) - len(new_docs)
            for i, embedding in enumerate(doc_embeddings):
                self.hnsw_index.add_items([embedding], [start_idx + i])

            self.hnsw_index.set_ef(self.ef)
            print(f"File {os.path.basename(file_path)} indexed successfully.")

    def embed_and_index(self):
        """Embed and index all loaded documents."""
        if not self.documents:
            print("No documents to embed and index.")
            return

        print(f"Embedding {len(self.documents)} document chunks.")
        texts = [doc.page_content for doc in self.documents]
        # Embedding with batch processing to avoid large prompts
        batch_size = 100  # Adjust as needed
        doc_embeddings = []
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i+batch_size]
            embeddings_batch = embeddings.embed_documents(batch_texts)
            doc_embeddings.extend(embeddings_batch)
            print(f"Embedded batch {i//batch_size +1}")

        doc_embeddings = np.array(doc_embeddings)

        # Add embeddings to HNSW index
        for i, embedding in enumerate(doc_embeddings):
            self.hnsw_index.add_items([embedding], [i])
            if i % 1000 == 0 and i > 0:
                print(f"Added {i} items to HNSW index.")

        self.hnsw_index.set_ef(self.ef)
        print("Documents added to HNSW index successfully.")

    def search_index(self, query, top_k=5):
        """Search the index for relevant documents based on the query."""
        query_embedding = embeddings.embed_documents([query])[0]
        labels, distances = self.hnsw_index.knn_query(np.array([query_embedding]), k=top_k)
        results = [{"text": self.documents[label].page_content, "distance": distance}
                   for label, distance in zip(labels[0], distances[0])]
        return results

    def generate_conversational_response(self, query):
        """Generate a conversational response based on the query and retrieved context."""
        results = self.search_index(query, top_k=5)

        if not results:
            return "I'm sorry, I couldn't find any relevant information to answer your question."

        # Construct context from retrieved documents
        context = "\n\n".join([f"{i+1}. {res['text']}" for i, res in enumerate(results)])

        # Create the prompt for ChatCompletion
        messages = [
            {"role": "system", "content": "You are a helpful and conversational assistant."},
            {"role": "user", "content": f"Use the following context to answer the question in a conversational manner.\n\nContext:\n{context}\n\nQuestion: {query}\nAnswer:"}
        ]

        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",  # Use "gpt-4" if you have access
                messages=messages,
                max_tokens=500,  # Limit the answer length
                temperature=0.7,
            )
            answer = response.choices[0].message['content'].strip()
            return answer
        except openai.error.InvalidRequestError as e:
            print(f"OpenAI API error: {e}")
            return "I'm sorry, I couldn't process your request at this time."
        except Exception as e:
            print(f"Error generating response: {e}")
            return "An error occurred while generating the response."

    def save_index_and_documents(self, index_file, documents_file):
        """Save the index and documents to disk."""
        self.hnsw_index.save_index(index_file)
        with open(documents_file, 'wb') as f:
            pickle.dump(self.documents, f)
        print("Index and documents saved to disk.")

    def load_index_and_documents(self, index_file, documents_file):
        """Load the index and documents from disk."""
        self.hnsw_index.load_index(index_file)
        with open(documents_file, 'rb') as f:
            self.documents = pickle.load(f)
        print("Index and documents loaded from disk.")

# ==========================
# Initialize Indexer and Index Documents
# ==========================

indexer = Indexer()
folder_path = UPLOAD_FOLDER

index_file = os.path.join(base_dir, 'index.bin')
documents_file = os.path.join(base_dir, 'documents.pkl')

if os.path.exists(index_file) and os.path.exists(documents_file):
    # Load existing index and documents
    indexer.load_index_and_documents(index_file, documents_file)
else:
    # Initialize and save index and documents
    indexer.load_documents_from_folder(folder_path)
    indexer.embed_and_index()
    indexer.save_index_and_documents(index_file, documents_file)

# ==========================
# Flask Routes
# ==========================

@app.route('/')
def home():
    return render_template('index.html')

# Route for file upload
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'files[]' not in request.files:
        return jsonify({"response": "No files part in the request."}), 400

    files = request.files.getlist('files[]')
    if not files or files[0].filename == '':
        return jsonify({"response": "No selected files."}), 400

    uploaded_files = []
    for file in files:
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            print(f"File {filename} uploaded successfully.")

            # After uploading, load and index the new file
            indexer.load_and_index_file(file_path)
            uploaded_files.append(filename)
        else:
            print(f"Invalid file type: {file.filename}")

    if uploaded_files:
        # Save the updated index and documents
        indexer.save_index_and_documents(index_file, documents_file)
        return jsonify({"response": f"Files {', '.join(uploaded_files)} uploaded and indexed successfully."}), 200
    else:
        return jsonify({"response": "No valid files were uploaded."}), 400

# Route to get the list of uploaded files
@app.route('/files', methods=['GET'])
def list_files():
    files = os.listdir(app.config['UPLOAD_FOLDER'])
    return jsonify({"files": files})

# Route for asking a question
@app.route('/ask', methods=['POST'])
def ask():
    data = request.get_json()
    question = data.get('question')
    if not question:
        return jsonify({"response": "No question provided!"}), 400

    answer = indexer.generate_conversational_response(question)
    if answer:
        return jsonify({"response": answer}), 200
    else:
        return jsonify({"response": "I'm sorry, I couldn't find any relevant information."}), 404

# Route to add a scenario
@app.route('/add_scenario', methods=['POST'])
def add_scenario():
    data = request.get_json()
    scenario_text = data.get('scenario')
    if not scenario_text:
        return jsonify({"response": "No scenario text provided!"}), 400

    # Save the scenario to a text file in the data folder
    scenario_filename = f"scenario_{int(time.time())}.txt"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], scenario_filename)
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(scenario_text)
        print(f"New scenario saved to {file_path}")

        # Load and index the new scenario file
        indexer.load_and_index_file(file_path)

        # Save the updated index and documents
        indexer.save_index_and_documents(index_file, documents_file)

        return jsonify({"response": "New scenario added successfully."}), 200
    except Exception as e:
        print(f"Error saving new scenario: {e}")
        return jsonify({"response": "Error adding new scenario."}), 500

# ==========================
# Run the Flask App
# ==========================

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
